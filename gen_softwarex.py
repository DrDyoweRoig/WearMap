"""Generate Wearmap_SoftwareX.docx — SoftwareX journal article."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page layout (A4, 2.5 cm margins)
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(2.5)
    s.left_margin = s.right_margin = Cm(2.5)

# ── Colour palette
DARK   = RGBColor(0x1A, 0x1A, 0x1A)
BLUE   = RGBColor(0x00, 0x56, 0x96)   # SoftwareX accent
GREY   = RGBColor(0x44, 0x44, 0x44)

# ── Helper functions
def para(text="", bold=False, italic=False, size=11, color=DARK,
         align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.name = "Times New Roman"
    return p

def mixed(parts, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    """parts: list of (text, bold, italic, size, color)"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    for text, bold, italic, size, color in parts:
        r = p.add_run(text)
        r.bold   = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.font.name = "Times New Roman"
    return p

def heading(text, level=1):
    sizes = {1: 14, 2: 12, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(sizes.get(level, 11))
    r.font.color.rgb = BLUE if level == 1 else DARK
    r.font.name = "Times New Roman"
    return p

def bullet(text, bold_prefix=None, indent=True):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    if bold_prefix and text.startswith(bold_prefix):
        rb = p.add_run(bold_prefix)
        rb.bold = True; rb.font.name = "Times New Roman"; rb.font.size = Pt(11)
        rb.font.color.rgb = DARK
        rr = p.add_run(text[len(bold_prefix):])
        rr.font.name = "Times New Roman"; rr.font.size = Pt(11)
        rr.font.color.rgb = DARK
    else:
        r = p.add_run(text)
        r.font.name = "Times New Roman"; r.font.size = Pt(11)
        r.font.color.rgb = DARK

def hline():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), "AAAAAA")
    pBdr.append(bot); pPr.append(pBdr)

def table2(rows_data, col_widths=(6, 10), header_row=True):
    """Two-column table. col_widths in cm."""
    from docx.shared import Cm as C
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    for i, (c1, c2) in enumerate(rows_data):
        row = tbl.add_row()
        for j, (cell, text) in enumerate([(row.cells[0], c1), (row.cells[1], c2)]):
            cell.width = C(col_widths[j])
            for rp in cell.paragraphs:
                rp.clear()
            rp = cell.paragraphs[0]
            r = rp.add_run(text)
            r.font.name = "Times New Roman"
            r.font.size  = Pt(10)
            r.font.color.rgb = DARK
            if header_row and i == 0:
                r.bold = True
                from docx.oxml import OxmlElement as OE
                from docx.oxml.ns import qn as Q
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OE("w:shd")
                shd.set(Q("w:val"), "clear")
                shd.set(Q("w:color"), "auto")
                shd.set(Q("w:fill"), "D6E4F0")
                tcPr.append(shd)
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
# TITLE & AUTHORS
# ════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(8)
r = p.add_run("Wearmap: An open-source Python application for quantitative dental wear image analysis and Puech directionality classification")
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = BLUE; r.font.name = "Times New Roman"

para("[Author names]", italic=True, size=11, color=GREY, space_after=2)
para("[Affiliations]", italic=True, size=10, color=GREY, space_after=2)
para("[Corresponding author e-mail]", italic=True, size=10, color=GREY, space_after=10)
hline()

# ════════════════════════════════════════════════════════════════
# ABSTRACT
# ════════════════════════════════════════════════════════════════
heading("Abstract", level=1)
para(
    "Wearmap is an open-source desktop application written in Python and PyQt6 that provides an "
    "integrated environment for quantitative dental wear image analysis. The software enables "
    "researchers to perform distance, area, and angle measurements on microscopic and macroscopic "
    "dental images, mark pit features with coordinate export, and apply real-world scale calibration "
    "using a wide range of physical units (from picometres to metres). A central feature of Wearmap "
    "is the automated implementation of the Puech (1983) tooth-wear directionality classification "
    "system, which categorises wear striations into four directions—Horizontal (P1), Vertical (P2), "
    "Mesiodistal (P3), and Distomesial (P4)—based on the measured slope angle and the lateral side "
    "of the tooth. Per-image measurement states are preserved during multi-image navigation, and "
    "results are exported to structured Excel workbooks containing both a detailed measurement sheet "
    "and a statistical summary sheet with Puech-category counts, mean distances, and standard "
    "deviations. Additional tools include non-destructive image filtering, image rotation and "
    "mirroring, and a visual scale bar overlay that is embedded directly into exported annotated "
    "images. Wearmap is distributed as a self-contained Windows executable and is designed as an "
    "open-source alternative to proprietary tools such as SigmaScan Pro, promoting reproducibility "
    "and methodological standardisation in paleoanthropological and forensic odontological research.",
    size=11, space_after=8
)

heading("Keywords", level=2)
para("dental wear • image analysis • Puech classification • Python • open-source • microscopy • paleoanthropology • PyQt6",
     italic=True, size=10, space_after=10)
hline()

# ════════════════════════════════════════════════════════════════
# CODE METADATA TABLE
# ════════════════════════════════════════════════════════════════
heading("Code Metadata", level=1)
table2([
    ("Field", "Value"),
    ("Current code version", "v1.0"),
    ("Permanent link to executables", "[URL — to be added upon repository publication]"),
    ("Legal Software License", "[To be defined by authors]"),
    ("Computing platform / OS", "Microsoft Windows 10 / 11 (64-bit)"),
    ("Installation requirements & dependencies",
     "No installation required; standalone executable provided.\n"
     "Source dependencies: Python ≥3.10, PyQt6 ≥6.4, Pillow ≥9.0, NumPy ≥1.21, openpyxl ≥3.0, pyinstaller ≥5.0"),
    ("Link to user manual", "Included with distribution (Wearmap_Manual.docx)"),
    ("Support email for questions", "[To be added]"),
], col_widths=[6.5, 9.5])
hline()

# ════════════════════════════════════════════════════════════════
# 1. MOTIVATION AND SIGNIFICANCE
# ════════════════════════════════════════════════════════════════
heading("1. Motivation and Significance", level=1)
para(
    "Dental wear analysis occupies a pivotal role in paleoanthropology, forensic odontology, and "
    "archaeological dietary reconstruction. The morphology and directionality of wear facets, "
    "striations, and pit features on tooth enamel encode information about diet, food-processing "
    "techniques, and tool use across human populations. Quantitative characterisation of these "
    "features requires measurement tools capable of operating on high-magnification images obtained "
    "by scanning electron microscopy (SEM), confocal microscopy, or macro photography.",
    size=11, space_after=6
)
para(
    "The most widely applied framework for classifying wear-striation directionality remains that of "
    "Puech (1983), which assigns each measured wear mark to one of four categories—Horizontal (P1), "
    "Vertical (P2), Mesiodistal (P3), or Distomesial (P4)—on the basis of the striation slope "
    "angle and the anatomical position of the tooth (jaw and side). Despite its widespread adoption, "
    "this classification has historically been implemented as post-hoc statistical scripts (typically "
    "in SPSS), requiring researchers to export raw angle measurements from image-analysis software, "
    "annotate each measurement with its tooth context, and run a separate classification routine. "
    "This multi-step workflow is error-prone, time-consuming, and difficult to reproduce across "
    "different research groups.",
    size=11, space_after=6
)
para(
    "Existing image-analysis solutions present significant limitations for dental wear research. "
    "SigmaScan Pro, long the de-facto standard in the field, is a commercial application that is "
    "no longer actively maintained and is unavailable on modern 64-bit operating systems. ImageJ/Fiji "
    "provides general-purpose measurement tools but lacks integrated Puech classification, real-world "
    "calibration to microscopy-relevant units (nanometres, micrometres), or structured statistical "
    "export. Purpose-built dental-wear plugins for ImageJ exist but typically address only a subset "
    "of the required workflow.",
    size=11, space_after=6
)
para(
    "Wearmap addresses these gaps by integrating, in a single open-source desktop application, all "
    "steps of the dental wear image-analysis pipeline: image loading and enhancement, measurement "
    "with real-world calibration, automated Puech classification based on tooth context, and "
    "statistical export. The application is distributed as a self-contained Windows executable "
    "requiring no Python installation, lowering the barrier to adoption for non-technical researchers. "
    "By releasing Wearmap as open-source software, we also aim to enable community-driven "
    "improvements and methodological transparency.",
    size=11, space_after=10
)

# ════════════════════════════════════════════════════════════════
# 2. SOFTWARE DESCRIPTION
# ════════════════════════════════════════════════════════════════
heading("2. Software Description", level=1)

heading("2.1 Architecture", level=2)
para(
    "Wearmap is structured following a Model–View–Controller (MVC) pattern implemented across "
    "four principal Python modules:",
    size=11, space_after=4
)
bullet("measurements.py — data models (Calibration, Measurement, MeasurementType) and the core compute_puech() function.", "measurements.py —")
bullet("canvas.py — QGraphicsView-based canvas that handles image display, tool interaction, overlay rendering (MeasurementItem), and annotated image export.", "canvas.py —")
bullet("main_window.py — QMainWindow with toolbar, side panel, measurement table, filter controls, and all export logic.", "main_window.py —")
bullet("dialogs.py — modal dialogs for scale calibration and application information.", "dialogs.py —")
para(
    "Measurement overlays are implemented as QGraphicsItem subclasses rendered with cosmetic pens, "
    "so line widths remain constant at all zoom levels. Per-image state (measurements, calibration, "
    "tooth context) is serialised in memory and restored transparently during folder navigation.",
    size=11, space_before=4, space_after=8
)

heading("2.2 Measurement Tools", level=2)
para("Wearmap provides six measurement tools, selectable from the toolbar or via keyboard shortcuts:", size=11, space_after=4)
bullet("Distance (D): two-click line measurement. Displays length, slope angle, and Puech class.", "Distance (D):")
bullet("Area (A): n-vertex polygon with Shoelace-formula area computation. Right-click or double-click to close.", "Area (A):")
bullet("Angle (G): three-point angle (vertex at second click).", "Angle (G):")
bullet("Count (C): click-to-mark counter with numbered circles; right-click to finalise.", "Count (C):")
bullet("Pits (P): click-to-place small-dot markers for pit features; coordinates exported per point.", "Pits (P):")
bullet("Calibrate (K): draw a line over a known reference feature, then enter its real-world length and unit.", "Calibrate (K):")
para(
    "All measurements are colour-coded, individually visible/hidden, and deletable from the side panel. "
    "A Labels toggle (L) hides text overlays without removing measurements, useful when capturing "
    "annotated images for publication.",
    size=11, space_before=4, space_after=8
)

heading("2.3 Puech Directionality Classification", level=2)
para(
    "The automated Puech classification is implemented in compute_puech(angle_deg, jaw, side). "
    "The signed slope angle θ ∈ [−90°, 90°] returned by slope_angle() "
    "(computed as arctan2(Δy, Δx) after normalising the line direction to left-to-right) "
    "is first converted to a normalised angle α ∈ [0°, 180°] following the "
    "convention of the original SPSS script: α = θ if θ ≥ 0, else α = 180 + θ.",
    size=11, space_after=6
)
para("Five angular zones are then evaluated (Table 1):", size=11, space_after=4)
table2([
    ("Zone (α, normalised)", "Puech class"),
    ("0° – 22.5° and 157.5° – 180°", "P1 — Horizontal"),
    ("22.5° – 67.5°  (\\  diagonal)", "P3 Mesiodistal if Lower-Left or Upper-Right;\nP4 Distomesial if Lower-Right or Upper-Left"),
    ("67.5° – 112.5°", "P2 — Vertical"),
    ("112.5° – 157.5°  (/  diagonal)", "P3 Mesiodistal if Lower-Right or Upper-Left;\nP4 Distomesial if Lower-Left or Upper-Right"),
], col_widths=[7, 9])
para("Table 1. Puech classification zones. For diagonal zones, both jaw and side jointly determine P3/P4: the effective mesiodistal direction depends on whether the tooth pair is Lower-Left / Upper-Right (mesial rightward in buccal view) or Lower-Right / Upper-Left (mesial leftward in buccal view).",
     italic=True, size=9, color=GREY, space_after=8)
para(
    "The tooth context (jaw, side, tooth type) is set via three drop-down menus in the side panel "
    "before measurement and propagates automatically to all Distance measurements in the current image. "
    "Changing the context updates all existing classifications in real time.",
    size=11, space_after=8
)

heading("2.4 Scale Calibration", level=2)
para(
    "Real-world calibration is performed by drawing a Distance line over a scale bar or reference "
    "feature of known length, then entering the known value and selecting a unit from a list spanning "
    "pm, nm, Å, μm, mm, cm, m, km, in, and ft. The resulting scale factor (real length / "
    "pixel length) is applied to all subsequent Distance and Area measurements. Calibration is stored "
    "per image and included in all exports.",
    size=11, space_after=8
)

heading("2.5 Export", level=2)
para(
    "Measurements are exported to Excel (.xlsx), CSV, or tab-separated text. The Excel workbook "
    "contains two sheets:",
    size=11, space_after=4
)
bullet("Detail: one row per measurement, with columns Image, ID, Type, Tooth, Jaw, Side, Unit, Value, Slope (°), Puech, Points (px).", "Detail:")
bullet("Summary: one row per Puech category (P1–P4) per image, reporting count, mean distance, and sample standard deviation; a bold TOTAL row aggregates all distance measurements per image.", "Summary:")
para(
    "For CSV/TXT export, two separate files are generated (base name and base name_summary). "
    "This structure allows direct import into statistical software for between-group comparisons.",
    size=11, space_before=4, space_after=8
)

heading("2.6 Image Enhancement", level=2)
para(
    "A non-destructive filter pipeline (implemented via Pillow) offers brightness, contrast, "
    "sharpness, and Gaussian blur adjustments (−1 – +100 range), togglable "
    "grayscale and colour inversion, and four special presets: Auto Levels, Histogram Equalisation, "
    "Edge Detection, and Sharpen. Filters operate on an in-memory copy of the original image; "
    "the source file is never modified.",
    size=11, space_after=10
)

# ════════════════════════════════════════════════════════════════
# 3. ILLUSTRATIVE EXAMPLE
# ════════════════════════════════════════════════════════════════
heading("3. Illustrative Example", level=1)
para(
    "A typical Wearmap workflow for the quantitative analysis of buccal wear striations on an "
    "upper-right first molar (UM1R) proceeds as follows.",
    size=11, space_after=6
)
para(
    "The researcher opens a SEM image (TIFF, 2048 × 1536 px) via File → Open Image or "
    "drag-and-drop. The Tooth Context panel is immediately configured: Jaw = Upper, Side = Right, "
    "Tooth = M1. The calibration tool (K) is then activated; the researcher draws a line over the "
    "SEM scale bar (e.g., 100 μm) and enters 100 in the calibration dialog with unit μm. "
    "The status bar confirms the resulting pixel scale.",
    size=11, space_after=6
)
para(
    "Five wear striations are measured using the Distance tool (D). For each striation, Wearmap "
    "displays the length in μm, the normalised slope angle, and the automatically assigned "
    "Puech category (e.g., “12.45 μm ↘ 34.2° P3”). If labels obscure "
    "underlying morphology, the Labels toggle (L) hides all text overlays while preserving measurements. "
    "Two additional pit features are marked with the Pits tool (P); their pixel coordinates will "
    "appear in the export.",
    size=11, space_after=6
)
para(
    "Export (Ctrl+E) generates measurements.xlsx. The Detail sheet lists all seven measurements with "
    "their Puech classes; the Summary sheet shows, for this image, three P3 (Mesiodistal) striations "
    "with mean 14.3 μm (SD 2.1 μm), two P4 (Distomesial) striations, and overall statistics "
    "in the TOTAL row. The annotated image (File → Save Image, Ctrl+Shift+S) is saved as a "
    "PNG at the original SEM resolution with all overlays rendered at native scale, suitable for "
    "direct inclusion in publications.",
    size=11, space_after=6
)
para(
    "When the researcher opens additional images from the same folder, the ◀ Prev / "
    "Next ▶ buttons navigate between them while preserving all measurements and the "
    "calibration for each image independently. The final Export aggregates all images into a single "
    "workbook.",
    size=11, space_after=6
)
para("[Figure 1: Screenshot of the Wearmap interface showing a SEM image of a molar with five distance measurements, the Tooth Context panel (Jaw=Upper, Side=Right, Tooth=M1), and the measurement table with Puech classification.]",
     italic=True, size=10, color=GREY, space_after=10)

# ════════════════════════════════════════════════════════════════
# 4. IMPACT
# ════════════════════════════════════════════════════════════════
heading("4. Impact", level=1)
para(
    "Wearmap directly addresses the reproducibility deficit in quantitative dental wear research. "
    "By encoding the Puech (1983) classification algorithm as an auditable open-source function "
    "with a documented test matrix (all jaw × side × angular zone combinations verified), "
    "the software eliminates ambiguity in how individual laboratories interpret boundary conditions "
    "and side-dependent directionality assignments.",
    size=11, space_after=6
)
para(
    "The integrated statistical summary (mean, standard deviation per Puech category per image) "
    "produces output directly compatible with between-population comparisons and meta-analyses "
    "without requiring post-processing in external statistical packages. Researchers who previously "
    "relied on a SigmaScan Pro → SPSS pipeline can now perform the complete analysis in a "
    "single session.",
    size=11, space_after=6
)
para(
    "Wearmap is applicable beyond paleoanthropology: any discipline requiring directional striation "
    "analysis on hard tissues (forensic odontology, tribology, materials characterisation) can "
    "benefit from its measurement and export capabilities. The calibration system supports the full "
    "range of microscopy scales, from nanometre-level atomic-force microscopy to millimetre-level "
    "macro photography.",
    size=11, space_after=6
)
para(
    "As a standalone executable, Wearmap requires no Python environment or package management, "
    "making it immediately usable by researchers without computational backgrounds. The open-source "
    "codebase allows adaptation to additional classification schemes or export formats as community "
    "needs evolve.",
    size=11, space_after=10
)

# ════════════════════════════════════════════════════════════════
# 5. CONCLUSIONS
# ════════════════════════════════════════════════════════════════
heading("5. Conclusions", level=1)
para(
    "Wearmap provides a complete, open-source solution for quantitative dental wear image analysis, "
    "combining flexible measurement tools, automated Puech (1983) directionality classification, "
    "real-world scale calibration, non-destructive image enhancement, and structured statistical "
    "export in a single desktop application. By eliminating the dependence on proprietary and "
    "discontinued software, and by making the classification algorithm fully transparent and "
    "verifiable, Wearmap contributes to the methodological standardisation and long-term "
    "reproducibility of dental wear research.",
    size=11, space_after=10
)

# ════════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENTS
# ════════════════════════════════════════════════════════════════
heading("Acknowledgements", level=1)
para("[To be completed by authors.]", italic=True, size=11, space_after=10)

# ════════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════════
heading("References", level=1)

refs = [
    "Puech, P.F. (1983). Tooth wear in La Ferrassie man. American Journal of Physical Anthropology, 61(2), 149–157. https://doi.org/10.1002/ajpa.1330610204",
    "Puech, P.F., Albertini, H., & Serratrice, C. (1983). Tooth microwear and dietary patterns in early hominids from Laetoli, Hadar and Olduvai. Journal of Human Evolution, 12(8), 721–729.",
    "Brothwell, D.R. (1981). Digging Up Bones (3rd ed.). British Museum (Natural History) / Cornell University Press.",
    "Scott, R.S., Ungar, P.S., Bergstrom, T.S., Brown, C.A., Grine, F.E., Teaford, M.F., & Walker, A. (2005). Dental microwear texture analysis shows within-species diet variability in fossil hominins. Nature, 436, 693–695. https://doi.org/10.1038/nature03822",
    "Schindelin, J., Arganda-Carreras, I., Frise, E., Kaynig, V., Longair, M., Pietzsch, T., ... & Cardona, A. (2012). Fiji: an open-source platform for biological-image analysis. Nature Methods, 9(7), 676–682. https://doi.org/10.1038/nmeth.2019",
    "Riverbank Computing. (2023). PyQt6 Reference Guide. https://www.riverbankcomputing.com/static/Docs/PyQt6/",
    "Clark Erskine-Smith, J. et al. (2024). Pillow (PIL Fork) Documentation. https://pillow.readthedocs.io/",
    "Harris, C.R., Millman, K.J., van der Walt, S.J. et al. (2020). Array programming with NumPy. Nature, 585, 357–362. https://doi.org/10.1038/s41586-020-2649-2",
    "The Pandas Development Team. (2023). openpyxl: A Python library to read/write Excel 2010 xlsx/xlsm files. https://openpyxl.readthedocs.io/",
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    r1 = p.add_run(f"[{i}] ")
    r1.bold = True; r1.font.size = Pt(10); r1.font.name = "Times New Roman"; r1.font.color.rgb = DARK
    r2 = p.add_run(ref)
    r2.font.size = Pt(10); r2.font.name = "Times New Roman"; r2.font.color.rgb = DARK

OUT = r"C:\Users\alber\Documents\code\openscan\Wearmap_SoftwareX.docx"
doc.save(OUT)
print(f"Saved: {OUT}")
