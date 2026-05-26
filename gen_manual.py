from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

GREEN_DARK  = RGBColor(0x14, 0x53, 0x2D)
GREEN_MED   = RGBColor(0x16, 0xA3, 0x4A)
GREY        = RGBColor(0x4B, 0x7A, 0x5A)
BLACK       = RGBColor(0x1A, 0x1A, 0x1A)

def fmt(run, bold=False, size=11, color=BLACK):
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    fmt(r, bold=True, size=16, color=GREEN_DARK)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), "16A34A")
    pBdr.append(bot)
    pPr.append(pBdr)

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    fmt(r, bold=True, size=12, color=GREEN_MED)

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    fmt(r, size=11)

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix and text.startswith(bold_prefix):
        rb = p.add_run(bold_prefix)
        fmt(rb, bold=True, size=11, color=GREEN_DARK)
        rr = p.add_run(text[len(bold_prefix):])
        fmt(rr, size=11)
    else:
        fmt(p.add_run(text), size=11)

# ── TITLE
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("WEARMAP")
r.bold = True; r.font.name = "Calibri"; r.font.size = Pt(36)
r.font.color.rgb = GREEN_DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Manual d\u2019usuari  \u00b7  Versi\u00f3 1.0")
r.font.name = "Calibri"; r.font.size = Pt(14); r.font.color.rgb = GREEN_MED

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Programari d\u2019an\u00e0lisi d\u2019imatge per a la recerca en desgast dental")
r.font.name = "Calibri"; r.font.size = Pt(11); r.font.color.rgb = GREY

doc.add_page_break()

# ── 1
h1("1. Introducci\u00f3")
body("Wearmap \u00e9s un programari d\u2019an\u00e0lisi d\u2019imatge dissenyat espec\u00edficament per a la recerca en desgast dental. Permet realitzar mesures de dist\u00e0ncia, \u00e0rea i angle sobre imatges microsc\u00f2piques, classificar el desgast seguint el model de Puech (1983), i exportar totes les dades a Excel o CSV.")
body("El programa no requereix instal\u00b7laci\u00f3: n\u2019hi ha prou amb executar el fitxer Wearmap.exe.")

# ── 2
h1("2. Requisits del sistema")
bullet("Sistema operatiu: Windows 10 o Windows 11 (64 bits)", "Sistema operatiu:")
bullet("Mem\u00f2ria RAM: m\u00ednim 4 GB", "Mem\u00f2ria RAM:")
bullet("Espai en disc: ~150 MB per a l\u2019executable", "Espai en disc:")
bullet("Resoluci\u00f3 de pantalla: m\u00ednim 1280 \u00d7 800 px", "Resoluci\u00f3 de pantalla:")

# ── 3
h1("3. Primers passos")
h2("3.1 Obrir una imatge")
body("Hi ha tres maneres d\u2019obrir una imatge:")
bullet("Feu clic al bot\u00f3 \U0001f4c2 Open de la barra d\u2019eines.")
bullet("Aneu al men\u00fa File \u2192 Open Image.")
bullet("Arrossegueu i deixeu anar (drag & drop) un fitxer d\u2019imatge directament sobre la finestra.")
body("Formats acceptats: PNG, JPG, JPEG, TIF, TIFF, BMP, GIF, WEBP.")

h2("3.2 Navegaci\u00f3 entre imatges")
body("Si la imatge oberta es troba dins d\u2019una carpeta amb m\u00e9s imatges, els botons \u25c0 Prev i Next \u25b6 de la barra d\u2019eines permeten navegar-hi. Cada imatge mant\u00e9 les seves pr\u00f2pies mesures i calibraci\u00f3 independentment.")

h2("3.3 Zoom i desplacement")
bullet("Roda del ratol\u00ed: zoom d\u2019entrada i sortida.")
bullet("Bot\u00f3 Select (S) actiu + arrossegar: desplacement per la imatge.")
bullet("\U0001f50d+ / \U0001f50d\u2212 / Fit: zoom manual i ajust a la finestra.")

# ── 4
h1("4. Context de la dent")
body("Abans de realitzar mesures de dist\u00e0ncia, cal configurar el context de la dent al panell lateral dret. Aquesta informaci\u00f3 s\u2019utilitza per calcular autom\u00e0ticament la classificaci\u00f3 de Puech.")
h2("4.1 Par\u00e0metres")
bullet("Tooth (dent): I1, I2, C, Pm3, Pm4, M1, M2, M3", "Tooth (dent):")
bullet("Jaw (arcada): Upper (maxil\u00b7la) o Lower (mand\u00edbula)", "Jaw (arcada):")
bullet("Side (costat): Right (dret) o Left (esquerre)", "Side (costat):")
body("Qualsevol canvi en aquests par\u00e0metres s\u2019aplica immediatament a totes les mesures de la imatge actual.")

# ── 5
h1("5. Eines de mesura")
body("Seleccioneu l\u2019eina des de la barra d\u2019eines o amb la tecla de drecera corresponent.")

h2("5.1 Select  [S]")
body("Mode de selecci\u00f3 i desplacement. Permet seleccionar mesures existents fent clic sobre elles.")

h2("5.2 Distance  [D]")
body("Mesura la dist\u00e0ncia entre dos punts. Feu clic al punt inicial i al punt final. El resultat mostra la dist\u00e0ncia, la pendent (angle en graus) i la classificaci\u00f3 de Puech (P1\u2013P4) autom\u00e0ticament.")
body("Classificaci\u00f3 de Puech:")
bullet("P1 \u2014 Horitzontal (angle \u2264 22,5\u00b0 o \u2265 157,5\u00b0)")
bullet("P2 \u2014 Vertical (angle entre 67,5\u00b0 i 112,5\u00b0)")
bullet("P3 \u2014 Mesial")
bullet("P4 \u2014 Distal")

h2("5.3 Area  [A]")
body("Dibux un pol\u00edgon per mesurar una \u00e0rea. Feu clic per afegir v\u00e8rtexs. Tanqueu el pol\u00edgon amb el bot\u00f3 dret del ratol\u00ed o fent doble clic. Cal un m\u00ednim de 3 punts.")

h2("5.4 Angle  [G]")
body("Mesura l\u2019angle format per tres punts. Feu clic al primer extrem, al v\u00e8rtex i al segon extrem.")

h2("5.5 Count  [C]")
body("Compta objectes. Feu clic per marcar cada objecte (apareix un cercle numerat). Tanqueu el recompte amb el bot\u00f3 dret del ratol\u00ed.")

h2("5.6 Pits  [P]")
body("Marca pits o punts d\u2019inter\u00e8s amb un punt petit. Funcionament id\u00e8ntic al Count per\u00f2 sense numeraci\u00f3. Les coordenades de cada pit s\u2019exporten a la columna Points (px).")

h2("5.7 Calibrate  [K]")
body("Estableix l\u2019escala real de la imatge. Dibuixeu una l\u00ednia sobre un element de mesura conegut, introduiu la seva longitud real i la unitat (pm, nm, \u00c5, \u03bcm, mm, cm\u2026). Un cop calibrada, totes les mesures es mostren en unitats reals.")

h2("5.8 Labels  [L]")
body("Bot\u00f3 de commutaci\u00f3 que mostra o amaga les etiquetes de text sobre la imatge. Les mesures segueixen actives.")

# ── 6
h1("6. Filtratge d\u2019imatge")
body("El panell lateral cont\u00e9 la pestanya Filters amb els controls seg\u00fcents:")
bullet("Brightness / Contrast / Sharpness: lliscadors de \u221a100 a +100.", "Brightness / Contrast / Sharpness:")
bullet("Blur: desenfocament gaussi\u00e0.", "Blur:")
bullet("Grayscale: converteix a escala de grisos.", "Grayscale:")
bullet("Invert: inverteix els colors.", "Invert:")
bullet("Special filters: Auto Levels, Equalize, Edge Detect, Sharpen.", "Special filters:")
bullet("Reset: restaura la imatge original.", "Reset:")
body("Els filtres s\u00f3n no destructius: no modifiquen el fitxer original.")

# ── 7
h1("7. Exportaci\u00f3 de dades")
body("Feu clic al bot\u00f3 \U0001f4be Export (o Ctrl+E). S\u2019obriri un di\u00e0leg per escollir el format i la ubicaci\u00f3.")

h2("7.1 Formats disponibles")
bullet("Excel (.xlsx): recomanat. Genera un sol fitxer amb dos fulls.", "Excel (.xlsx):")
bullet("CSV (.csv): genera dos fitxers separats (_detail i _summary).", "CSV (.csv):")
bullet("Tab-separated (.txt): igual que CSV per\u00f2 separat per tabuladors.", "Tab-separated (.txt):")

h2("7.2 Full \u00abDetail\u00bb \u2014 dades individuals")
body("Una fila per cada mesura amb les columnes: Image, ID, Type, Tooth, Jaw, Side, Unit, Value, Slope (\u00b0), Puech, Points (px).")

h2("7.3 Full \u00abSummary\u00bb \u2014 estad\u00edstica per Puech")
body("Una fila per cada categoria de Puech (1\u20134) per imatge, m\u00e9s una fila TOTAL al final de cada imatge amb el recompte global, la mitjana i la desviaci\u00f3 est\u00e0ndard de totes les dist\u00e0ncies.")

# ── 8
h1("8. Dreceres de teclat")
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].paragraphs[0].add_run("Tecla").bold = True
hdr[1].paragraphs[0].add_run("Acci\u00f3").bold = True
shortcuts = [
    ("S", "Eina Select"),
    ("D", "Eina Distance"),
    ("A", "Eina Area"),
    ("G", "Eina Angle"),
    ("C", "Eina Count"),
    ("P", "Eina Pits"),
    ("K", "Eina Calibrate"),
    ("L", "Mostrar / amagar etiquetes"),
    ("Ctrl+O", "Obrir imatge"),
    ("Ctrl+E", "Exportar mesures"),
    ("Delete", "Eliminar mesura seleccionada"),
    ("\u2190 \u2192", "Imatge anterior / seg\u00fcent"),
    ("Ctrl + =", "Zoom d\u2019entrada"),
    ("Ctrl + -", "Zoom de sortida"),
    ("Ctrl + 0", "Ajustar a la finestra"),
]
for k, d in shortcuts:
    row = table.add_row()
    rb = row.cells[0].paragraphs[0].add_run(k)
    fmt(rb, bold=True, size=11, color=GREEN_DARK)
    fmt(row.cells[1].paragraphs[0].add_run(d), size=11)

# ── 9
h1("9. Resoluci\u00f3 de problemes")
h2("La imatge no s\u2019obre")
body("Verifiqueu que el format sigui compatible. Els fitxers TIFF multipla pots trigar uns segons.")
h2("Les etiquetes cobreixen la imatge")
body("Premeu L o feu clic al bot\u00f3 Labels per amagar-les temporalment.")
h2("El full Summary apareix buit")
body("El full Summary nom\u00e9s inclou mesures de tipus Distance amb Puech v\u00e0lid (1\u20134). Comproveu que Jaw i Side estiguin configurats.")
h2("El valor de Puech \u00e9s sempre \u00ab-\u00bb")
body("Configureu els desplegables Jaw i Side al panell Tooth Context abans de dibuixar les dist\u00e0ncies.")

# ── FINAL
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("Wearmap v1.0  \u00b7  Programari d\u2019an\u00e0lisi de desgast dental")
r.font.name = "Calibri"; r.font.size = Pt(10); r.font.color.rgb = GREY

doc.save("Wearmap_Manual.docx")
print("Done")
